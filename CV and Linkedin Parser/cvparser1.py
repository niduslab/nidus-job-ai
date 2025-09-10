import os
import re
import json
from datetime import datetime
import docx
import pdfplumber

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF and save to .txt file."""
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        # Save extracted text to local .txt file
        if text:
            txt_path = pdf_path.rsplit('.', 1)[0] + '_extracted.txt'
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(text)
            print(f"✅ Text extracted and saved to: {txt_path}")
        
        return text
    except Exception as e:
        print(f"❌ Error extracting PDF: {e}")
        return ""

def extract_text_from_docx(docx_path):
    """Extract text from DOCX file."""
    try:
        doc = docx.Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text
        
        return text
    except Exception as e:
        print(f"❌ Error extracting DOCX: {e}")
        return ""

def clean_text(text):
    """Clean and normalize text."""
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters but keep essential ones
    text = re.sub(r'[^\w\s@.,-/()\n]', ' ', text)
    return text.strip()

def extract_contact_info(text):
    """Extract comprehensive contact information."""
    contact_info = {}
    
    # Name (first line or line without email/phone)
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    for line in lines[:5]:  # Check first 5 lines
        if not re.search(r'[@+\d]', line) and len(line.split()) <= 5:
            contact_info['name'] = line.title()
            break
    
    # Email (enhanced pattern)
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    email_match = re.search(email_pattern, text)
    contact_info['email'] = email_match.group(0) if email_match else ""
    
    # Phone (comprehensive pattern)
    phone_patterns = [
        r'\+?\d{1,3}[-.\s]?\(?\d{3,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}',
        r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
        r'\d{10,15}'
    ]
    
    for pattern in phone_patterns:
        phone_match = re.search(pattern, text)
        if phone_match:
            phone = phone_match.group(0)
            # Clean phone number
            phone = re.sub(r'[^\d+]', '', phone)
            if len(phone) >= 10:
                contact_info['phone'] = phone
                break
    
    # Location (city, state, country)
    location_pattern = r'(?:^|\n).*?([A-Z][a-z]+,?\s*[A-Z][a-z]*,?\s*[A-Z][a-z]*)'
    location_match = re.search(location_pattern, text)
    contact_info['location'] = location_match.group(1).strip() if location_match else ""
    
    # LinkedIn
    linkedin_patterns = [
        r'linkedin\.com/in/[\w-]+',
        r'linkedin\.com/[\w-]+',
        r'in/[\w-]+'
    ]
    
    for pattern in linkedin_patterns:
        linkedin_match = re.search(pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact_info['linkedin'] = linkedin_match.group(0)
            break
    
    # GitHub
    github_pattern = r'github\.com/[\w-]+'
    github_match = re.search(github_pattern, text, re.IGNORECASE)
    contact_info['github'] = github_match.group(0) if github_match else ""
    
    return contact_info

def extract_summary_objective(text):
    """Extract professional summary or objective."""
    summary_patterns = [
        r'(summary|professional summary|career summary|profile|objective|career objective)[:\n](.*?)(?=\n\s*[A-Z][a-z]+|\n\s*\n|\Z)',
        r'(about me|about|introduction)[:\n](.*?)(?=\n\s*[A-Z][a-z]+|\n\s*\n|\Z)'
    ]
    
    for pattern in summary_patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            summary = match.group(2).strip()
            summary = re.sub(r'\s+', ' ', summary)
            if len(summary) > 50:  # Reasonable summary length
                return summary[:500]  # Limit length
    
    return ""

def extract_skills(text):
    """Extract technical and soft skills comprehensively."""
    skills = {
        'technical_skills': [],
        'soft_skills': [],
        'all_skills': []
    }
    
    # Technical skills database
    technical_skills_db = [
        # Programming Languages
        'python', 'java', 'javascript', 'c++', 'c#', 'c', 'r', 'php', 'ruby', 'go', 
        'swift', 'kotlin', 'scala', 'rust', 'typescript', 'dart', 'matlab', 'sql',
        
        # Web Technologies
        'html', 'css', 'react', 'angular', 'vue.js', 'node.js', 'express.js', 
        'django', 'flask', 'spring boot', 'asp.net', 'laravel', 'bootstrap', 'jquery',
        
        # Databases
        'mysql', 'postgresql', 'mongodb', 'sqlite', 'oracle', 'sql server', 
        'redis', 'elasticsearch', 'cassandra', 'dynamodb',
        
        # Cloud & DevOps
        'aws', 'azure', 'google cloud', 'docker', 'kubernetes', 'jenkins', 
        'git', 'github', 'gitlab', 'terraform', 'ansible', 'linux', 'unix',
        
        # Data Science & AI
        'machine learning', 'deep learning', 'data science', 'pandas', 'numpy', 
        'matplotlib', 'seaborn', 'scikit-learn', 'tensorflow', 'pytorch', 'keras',
        'tableau', 'power bi', 'excel', 'jupyter', 'spark', 'hadoop',
        
        # Mobile Development
        'android', 'ios', 'react native', 'flutter', 'xamarin',
        
        # Other Technologies
        'microservices', 'restful api', 'graphql', 'websockets', 'oauth', 'jwt'
    ]
    
    # Soft skills database
    soft_skills_db = [
        'leadership', 'communication', 'teamwork', 'problem solving', 
        'project management', 'time management', 'analytical thinking', 
        'creativity', 'adaptability', 'critical thinking', 'collaboration',
        'public speaking', 'negotiation', 'mentoring', 'strategic planning'
    ]
    
    # Find skills section
    skills_patterns = [
        r'(skills|technical skills|key skills|core competencies|technologies)[:\n](.*?)(?=\n\s*[A-Z][a-z]+|\n\s*\n|\Z)',
        r'(programming languages|tools|frameworks)[:\n](.*?)(?=\n\s*[A-Z][a-z]+|\n\s*\n|\Z)'
    ]
    
    skills_text = ""
    for pattern in skills_patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            skills_text += " " + match.group(2)
    
    # If no specific skills section, search entire text
    if not skills_text.strip():
        skills_text = text
    
    # Extract technical skills
    text_lower = skills_text.lower()
    for skill in technical_skills_db:
        if re.search(r'\b' + re.escape(skill.lower()) + r'\b', text_lower):
            skills['technical_skills'].append(skill.title())
    
    # Extract soft skills
    for skill in soft_skills_db:
        if re.search(r'\b' + re.escape(skill.lower()) + r'\b', text_lower):
            skills['soft_skills'].append(skill.title())
    
    # Remove duplicates
    skills['technical_skills'] = list(set(skills['technical_skills']))
    skills['soft_skills'] = list(set(skills['soft_skills']))
    skills['all_skills'] = skills['technical_skills'] + skills['soft_skills']
    
    return skills

def extract_education(text):
    """Extract comprehensive education information."""
    education = []
    
    # Find education section
    education_patterns = [
        r'(education|academic|qualification|degree)[:\n](.*?)(?=\n\s*[A-Z][a-z]+|\n\s*\n|\Z)',
        r'(university|college|institute)[:\n](.*?)(?=\n\s*[A-Z][a-z]+|\n\s*\n|\Z)'
    ]
    
    education_text = ""
    for pattern in education_patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            education_text = match.group(2)
            break
    
    if not education_text:
        education_text = text  # Search entire text if no section found
    
    # Degree patterns
    degree_patterns = [
        r'(bachelor|master|phd|doctorate|b\.?sc|m\.?sc|b\.?tech|m\.?tech|mba|bba|diploma|certificate|b\.?e|m\.?e|b\.?a|m\.?a)\.?\s*(of|in|-)?\s*([a-z\s]+)',
        r'(undergraduate|graduate|postgraduate)\s+(degree|program)\s*in\s*([a-z\s]+)'
    ]
    
    lines = education_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if len(line) < 10:  # Skip very short lines
            continue
            
        education_entry = {}
        
        # Extract degree
        for pattern in degree_patterns:
            match = re.search(pattern, line, re.IGNORECASE)
            if match:
                degree_type = match.group(1)
                field_of_study = match.group(3) if len(match.groups()) >= 3 else ""
                
                education_entry['degree'] = degree_type.title()
                education_entry['field_of_study'] = field_of_study.strip().title()
                break
        
        # Extract institution
        institution_pattern = r'(university|college|institute|school)\s+of\s+[\w\s]+|[\w\s]+(university|college|institute|school)'
        institution_match = re.search(institution_pattern, line, re.IGNORECASE)
        if institution_match:
            education_entry['institution'] = institution_match.group(0).title()
        
        # Extract year
        year_pattern = r'\b(19|20)\d{2}\b'
        year_matches = re.findall(year_pattern, line)
        if year_matches:
            education_entry['graduation_year'] = year_matches[-1]  # Take the latest year
        
        # Extract GPA/Score
        gpa_pattern = r'(\d+\.?\d*)\s*(gpa|cgpa|percentage|%)'
        gpa_match = re.search(gpa_pattern, line, re.IGNORECASE)
        if gpa_match:
            education_entry['gpa'] = gpa_match.group(1) + " " + gpa_match.group(2).upper()
        
        # Only add if we found degree information
        if education_entry and 'degree' in education_entry:
            education.append(education_entry)
    
    return education

def extract_experience(text):
    """Extract comprehensive work experience."""
    experience = []
    
    # Find experience section
    exp_patterns = [
        r'(experience|employment|career|professional experience|work experience)[:\n](.*?)(?=\n\s*[A-Z][a-z]+|\n\s*\n|\Z)',
        r'(work history|employment history)[:\n](.*?)(?=\n\s*[A-Z][a-z]+|\n\s*\n|\Z)'
    ]
    
    exp_text = ""
    for pattern in exp_patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            exp_text = match.group(2)
            break
    
    # If no section found, look for job title patterns in entire text
    if not exp_text:
        lines = text.split('\n')
        job_indicators = ['engineer', 'developer', 'analyst', 'manager', 'intern', 'specialist', 'consultant', 'lead', 'senior', 'junior']
        for line in lines:
            if any(indicator in line.lower() for indicator in job_indicators):
                exp_text += line + "\n"
    
    # Split into job entries (assuming jobs are separated by blank lines or dates)
    job_blocks = re.split(r'\n\s*\n|\n(?=\d{4})', exp_text)
    
    for block in job_blocks:
        if len(block.strip()) < 20:  # Skip very short blocks
            continue
            
        exp_entry = {}
        lines = [line.strip() for line in block.split('\n') if line.strip()]
        
        # Extract job title (usually first line or line with job keywords)
        job_title_patterns = [
            r'(senior|junior|lead|principal|chief)?\s*(software engineer|data scientist|product manager|business analyst|web developer|full stack developer|backend developer|frontend developer|devops engineer|qa engineer|intern)',
            r'(manager|director|coordinator|specialist|consultant|analyst|developer|engineer)'
        ]
        
        for line in lines:
            for pattern in job_title_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    exp_entry['position'] = line.title()
                    break
            if 'position' in exp_entry:
                break
        
        # Extract company name (look for "at Company", "Company Name", etc.)
        company_patterns = [
            r'at\s+([\w\s&.,]+?)(?:\s|$|\n)',
            r'([\w\s&.,]+?)\s+(inc|corp|llc|ltd|company|technologies|solutions)',
            r'^([\w\s&.,]+)$'  # Line that might be just company name
        ]
        
        for line in lines:
            for pattern in company_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match and len(match.group(1).strip()) > 2:
                    potential_company = match.group(1).strip()
                    # Avoid job titles as company names
                    if not re.search(r'(engineer|developer|manager|analyst|intern)', potential_company, re.IGNORECASE):
                        exp_entry['company'] = potential_company.title()
                        break
            if 'company' in exp_entry:
                break
        
        # Extract duration/dates
        date_patterns = [
            r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{4}\s*[-–]\s*(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{4}',
            r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{4}\s*[-–]\s*present',
            r'\d{4}\s*[-–]\s*\d{4}',
            r'\d{4}\s*[-–]\s*present',
            r'(\d+)\s*(years?|months?|yrs?)'
        ]
        
        duration_text = " ".join(lines)
        for pattern in date_patterns:
            match = re.search(pattern, duration_text, re.IGNORECASE)
            if match:
                exp_entry['duration'] = match.group(0)
                break
        
        # Extract description (remaining text)
        description_lines = []
        for line in lines:
            # Skip lines that are job title, company, or duration
            if ('position' in exp_entry and exp_entry['position'].lower() not in line.lower() and
                'company' in exp_entry and exp_entry['company'].lower() not in line.lower() and
                'duration' in exp_entry and exp_entry['duration'].lower() not in line.lower()):
                if len(line) > 20:  # Meaningful description
                    description_lines.append(line)
        
        if description_lines:
            exp_entry['description'] = " ".join(description_lines)
        
        # Only add if we have meaningful information
        if exp_entry and ('position' in exp_entry or 'company' in exp_entry):
            experience.append(exp_entry)
    
    return experience

def extract_projects(text):
    """Extract project information."""
    projects = []
    
    # Find projects section
    project_patterns = [
        r'(projects|portfolio|work samples|personal projects)[:\n](.*?)(?=\n\s*[A-Z][a-z]+|\n\s*\n|\Z)',
        r'(key projects|major projects|notable projects)[:\n](.*?)(?=\n\s*[A-Z][a-z]+|\n\s*\n|\Z)'
    ]
    
    project_text = ""
    for pattern in project_patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            project_text = match.group(2)
            break
    
    if project_text:
        # Split into individual projects
        project_blocks = re.split(r'\n\s*\n|\n(?=[A-Z][a-z])', project_text)
        
        for block in project_blocks:
            if len(block.strip()) > 30:  # Reasonable project description
                project_entry = {}
                lines = [line.strip() for line in block.split('\n') if line.strip()]
                
                # Project name (usually first line)
                if lines:
                    project_entry['name'] = lines[0]
                
                # Extract technologies
                tech_keywords = ['python', 'java', 'react', 'node.js', 'sql', 'html', 'css', 'javascript', 'django', 'flask', 'mongodb', 'mysql']
                technologies = []
                
                for line in lines:
                    for tech in tech_keywords:
                        if tech.lower() in line.lower():
                            technologies.append(tech)
                
                project_entry['technologies'] = list(set(technologies))
                project_entry['description'] = " ".join(lines[1:]) if len(lines) > 1 else ""
                
                # Extract duration if mentioned
                duration_pattern = r'(\d{4}|\d+\s*(months?|weeks?|days?))'
                duration_match = re.search(duration_pattern, block)
                if duration_match:
                    project_entry['duration'] = duration_match.group(0)
                
                projects.append(project_entry)
    
    return projects

def extract_certifications(text):
    """Extract certifications and licenses."""
    certifications = []
    
    # Find certifications section
    cert_patterns = [
        r'(certifications?|certificates?|licenses?|credentials?)[:\n](.*?)(?=\n\s*[A-Z][a-z]+|\n\s*\n|\Z)',
        r'(professional certifications?|industry certifications?)[:\n](.*?)(?=\n\s*[A-Z][a-z]+|\n\s*\n|\Z)'
    ]
    
    cert_text = ""
    for pattern in cert_patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            cert_text = match.group(2)
            break
    
    # Common certification patterns
    cert_keywords = [
        'aws certified', 'microsoft certified', 'google certified', 'cisco certified',
        'pmp', 'scrum master', 'six sigma', 'itil', 'comptia', 'cissp',
        'certified', 'certification', 'license'
    ]
    
    search_text = cert_text if cert_text else text
    lines = search_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if any(keyword in line.lower() for keyword in cert_keywords) and len(line) > 5:
            # Extract year if present
            year_match = re.search(r'\b(19|20)\d{2}\b', line)
            cert_info = {
                'name': line,
                'year': year_match.group(0) if year_match else ""
            }
            certifications.append(cert_info)
    
    return certifications

def calculate_experience_years(experience_list):
    """Calculate total years of experience."""
    if not experience_list:
        return "0 years"
    
    total_months = 0
    current_year = datetime.now().year
    
    for exp in experience_list:
        duration = exp.get('duration', '').lower()
        
        if 'present' in duration or 'current' in duration:
            # Extract start year
            year_match = re.search(r'\b(19|20)\d{2}\b', duration)
            if year_match:
                start_year = int(year_match.group(0))
                total_months += (current_year - start_year) * 12
        else:
            # Extract year range
            years = re.findall(r'\b(19|20)\d{2}\b', duration)
            if len(years) >= 2:
                start_year = int(years[0])
                end_year = int(years[-1])
                total_months += (end_year - start_year) * 12
            elif len(years) == 1:
                total_months += 12  # Assume 1 year
            
            # Check for month/year patterns
            if re.search(r'\d+\s*years?', duration):
                year_match = re.search(r'(\d+)\s*years?', duration)
                if year_match:
                    total_months += int(year_match.group(1)) * 12
            
            if re.search(r'\d+\s*months?', duration):
                month_match = re.search(r'(\d+)\s*months?', duration)
                if month_match:
                    total_months += int(month_match.group(1))
    
    total_years = max(1, total_months // 12) if total_months > 0 else 0
    
    if total_years == 0:
        return "Fresh Graduate"
    elif total_years == 1:
        return "1 year"
    else:
        return f"{total_years} years"

def parse_cv_text(text):
    """Parse CV text comprehensively with ATS-friendly sections."""
    
    # Clean text
    text = clean_text(text)
    
    # Extract all information
    contact_info = extract_contact_info(text)
    summary = extract_summary_objective(text)
    skills = extract_skills(text)
    education = extract_education(text)
    experience = extract_experience(text)
    projects = extract_projects(text)
    certifications = extract_certifications(text)
    
    # Calculate experience years
    experience_years = calculate_experience_years(experience)
    
    # Compile comprehensive data
    cv_data = {
        'contact_information': contact_info,
        'professional_summary': summary,
        'skills': skills,
        'education': education,
        'work_experience': experience,
        'projects': projects,
        'certifications': certifications,
        'years_of_experience': experience_years,
        'ats_score': calculate_ats_score(contact_info, skills, education, experience),
        'metadata': {
            'total_sections': sum([
                1 if contact_info.get('name') else 0,
                1 if summary else 0,
                1 if skills['all_skills'] else 0,
                1 if education else 0,
                1 if experience else 0,
                1 if projects else 0,
                1 if certifications else 0
            ]),
            'processed_at': datetime.now().isoformat(),
            'text_length': len(text)
        }
    }
    
    return cv_data

def calculate_ats_score(contact_info, skills, education, experience):
    """Calculate ATS friendliness score."""
    score = 0
    max_score = 100
    
    # Contact information (20 points)
    if contact_info.get('name'): score += 5
    if contact_info.get('email'): score += 5
    if contact_info.get('phone'): score += 5
    if contact_info.get('location'): score += 5
    
    # Skills (30 points)
    if skills['technical_skills']: score += 15
    if skills['soft_skills']: score += 10
    if len(skills['all_skills']) > 10: score += 5
    
    # Education (20 points)
    if education: score += 20
    
    # Experience (30 points)
    if experience: score += 20
    if len(experience) > 1: score += 5
    if any('description' in exp for exp in experience): score += 5
    
    return min(score, max_score)

def parse_cv_file(file_path):
    """Parse CV file and return comprehensive data."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    print(f"🔍 Processing: {file_path}")

    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif ext == ".docx":
        text = extract_text_from_docx(file_path)
    elif ext == ".doc":
        try:
            import doc2text
            doc = doc2text.Document(lang="eng")
            doc.read(str(file_path))
            doc.process()
            doc.extract_text()
            text = doc.get_text()
        except Exception as e:
            print(f"❌ Error extracting DOC with doc2text: {e}")
            text = ""
    else:
        raise ValueError("Unsupported file type. Only PDF, DOCX, and DOC are supported.")

    if not text.strip():
        raise ValueError("No text could be extracted from the file.")

    print(f"📄 Extracted {len(text)} characters")

    # Parse the text
    parsed_data = parse_cv_text(text)

    # Save results
    output_file = f"{os.path.splitext(file_path)[0]}_parsed_comprehensive.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(parsed_data, f, indent=2, ensure_ascii=False)

    print(f"💾 Results saved to: {output_file}")

    return parsed_data

def print_parsed_results(data):
    """Print parsed results in a formatted way."""
    print("\n" + "="*60)
    print("🎯 COMPREHENSIVE CV PARSING RESULTS")
    print("="*60)
    
    # Contact Information
    contact = data['contact_information']
    print(f"📋 CONTACT INFORMATION:")
    print(f"   Name: {contact.get('name', 'Not found')}")
    print(f"   Email: {contact.get('email', 'Not found')}")
    print(f"   Phone: {contact.get('phone', 'Not found')}")
    print(f"   Location: {contact.get('location', 'Not found')}")
    print(f"   LinkedIn: {contact.get('linkedin', 'Not found')}")
    print(f"   GitHub: {contact.get('github', 'Not found')}")
    
    # Summary
    if data['professional_summary']:
        print(f"\n📝 PROFESSIONAL SUMMARY:")
        print(f"   {data['professional_summary'][:200]}...")
    
    # Skills
    skills = data['skills']
    print(f"\n🔧 SKILLS:")
    print(f"   Technical Skills ({len(skills['technical_skills'])}): {', '.join(skills['technical_skills'][:10])}")
    if len(skills['technical_skills']) > 10:
        print(f"      ... and {len(skills['technical_skills']) - 10} more")
    
    print(f"   Soft Skills ({len(skills['soft_skills'])}): {', '.join(skills['soft_skills'])}")
    
    # Education
    print(f"\n🎓 EDUCATION ({len(data['education'])} entries):")
    for i, edu in enumerate(data['education'], 1):
        print(f"   {i}. {edu.get('degree', 'N/A')} in {edu.get('field_of_study', 'N/A')}")
        if edu.get('institution'):
            print(f"      Institution: {edu['institution']}")
        if edu.get('graduation_year'):
            print(f"      Year: {edu['graduation_year']}")
        if edu.get('gpa'):
            print(f"      GPA: {edu['gpa']}")
    
    # Work Experience
    print(f"\n💼 WORK EXPERIENCE ({len(data['work_experience'])} entries):")
    for i, exp in enumerate(data['work_experience'], 1):
        print(f"   {i}. {exp.get('position', 'N/A')} at {exp.get('company', 'N/A')}")
        if exp.get('duration'):
            print(f"      Duration: {exp['duration']}")
        if exp.get('description'):
            print(f"      Description: {exp['description'][:150]}...")
    
    # Projects
    if data['projects']:
        print(f"\n🚀 PROJECTS ({len(data['projects'])} found):")
        for i, proj in enumerate(data['projects'], 1):
            print(f"   {i}. {proj.get('name', 'N/A')}")
            if proj.get('technologies'):
                print(f"      Technologies: {', '.join(proj['technologies'])}")
            if proj.get('duration'):
                print(f"      Duration: {proj['duration']}")
    
    # Certifications
    if data['certifications']:
        print(f"\n🏆 CERTIFICATIONS ({len(data['certifications'])} found):")
        for i, cert in enumerate(data['certifications'], 1):
            print(f"   {i}. {cert.get('name', 'N/A')}")
            if cert.get('year'):
                print(f"      Year: {cert['year']}")
    
    # Summary Stats
    print(f"\n📊 SUMMARY STATISTICS:")
    print(f"   Years of Experience: {data['years_of_experience']}")
    print(f"   ATS Score: {data['ats_score']}/100")
    print(f"   Total Sections: {data['metadata']['total_sections']}")
    print(f"   Text Length: {data['metadata']['text_length']} characters")
    
    print("="*60)

if __name__ == "__main__":
    # Example usage
    file_path = "demoresume.pdf"  # Change this to your CV file
    
    try:
        parsed_data = parse_cv_file(file_path)
        print_parsed_results(parsed_data)
        
    except Exception as e:
        print(f"❌ Error processing CV: {e}")